import streamlit as st
import requests
import time
from datetime import datetime
import os
from dotenv import load_dotenv
from streamlit.components.v1 import html

# ========== 环境变量配置 ==========
load_dotenv()
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
headers = {
    "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
    "Content-Type": "application/json"
}

# ========== AI行为限制提示词 ========== 【新增部分】
SYSTEM_PROMPTS = {
    1: """请你严格遵守以下限制：
1. 仅回答用户当前明确提出的问题，不主动提供额外建议或提问
2. 避免包含"是否需要更多帮助？"等引导性语句
3. 禁止添加任何表情符号
4. 使用正式书面语气，禁止拟人化行为和第一人称代词""",

    2: """请你严格遵守以下限制：
1. 每次回答后主动提出3个与当前任务相关的后续问题
2. 用户问题模糊时，必须请求澄清
3. 每段回答后添加1个相关表情符号（如💡、🤔等）
4. 使用第一人称（如"我建议"）和口语化语气
5. 每次回答需包含对用户的积极认可（如"很好的问题！"）"""
}


# ========== 工具函数 ==========
def format_duration(seconds):
    """格式化时间显示"""
    if seconds <= 0:
        return "0秒"
    minutes, seconds = divmod(int(seconds), 60)
    hours, minutes = divmod(minutes, 60)
    return f"{hours:02}:{minutes:02}:{seconds:02}"


# ========== 初始化 Session 状态 ==========
required_states = {
    "messages": [],
    "research_notes": "",
    "user_message_timestamps": [],
    "ai_response_timestamps": [],
    "total_api_latency": 0.0,
    "should_redirect": False,
    "identity_confirmed": False,
    "total_duration": 0.0,
    "net_duration": 0.0
}

for key, value in required_states.items():
    if key not in st.session_state:
        st.session_state[key] = value

# ========== 页面配置 ==========
st.set_page_config(layout="wide", page_icon="🧪")
st.title("心理实验平台 🌱")

# ========== 侧边栏 ==========
with st.sidebar:
    st.subheader("被试人员身份")
    col1, col2 = st.columns([3, 2])
    with col1:
        user_id = st.radio(
            "请选择您的身份：",
            options=[1, 2],
            key="user_id",
            horizontal=True,
            on_change=lambda: st.session_state.update(identity_confirmed=False)
        )
    with col2:
        if st.button("🔒 确认身份",
                     help="请先选择身份后点击确认",
                     type="primary" if not st.session_state.identity_confirmed else "secondary"):
            if st.session_state.user_id in [1, 2]:
                st.session_state.identity_confirmed = True
                st.session_state.messages = []  # 切换身份时重置对话
                st.toast("身份验证成功，请开始实验！", icon="✅")
            else:
                st.error("请选择有效身份（1或2）")

    st.header("实验数据监控")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("有效对话轮次", f"{len(st.session_state.ai_response_timestamps)}次")
    with col2:
        st.metric("净交流时长", format_duration(st.session_state.net_duration))
    st.metric("总实验时长",
              format_duration(st.session_state.total_duration) if st.session_state.total_duration > 0 else "N/A")

    st.divider()
    if st.button("📥 导出对话记录"):
        if st.session_state.messages:
            from docx import Document

            doc = Document()
            doc.add_heading("实验对话记录", 0)
            doc.add_paragraph(f"实验总时长：{format_duration(st.session_state.total_duration)}")
            doc.add_paragraph(f"净交流时长：{format_duration(st.session_state.net_duration)}\n")

            for msg in st.session_state.messages:
                role = "用户" if msg["role"] == "user" else "AI助手"
                doc.add_paragraph(f"{role}：{msg['content']}")

            doc.save("dialogue_record.docx")
            with open("dialogue_record.docx", "rb") as f:
                st.download_button("下载文档", f, "dialogue_record.docx")
        else:
            st.warning("没有对话记录可导出")

# ========== 主界面 ==========
with st.container(height=600):
    st.subheader("AI 对话助手")

    if not st.session_state.identity_confirmed:
        st.info("👋 请先在左侧边栏选择被试身份并点击【确认身份】按钮")
        st.stop()

    # 实时显示对话记录
    for msg in st.session_state.messages:
        # 过滤系统提示不显示 【新增】
        if msg["role"] != "system":
            st.chat_message(msg["role"]).write(msg["content"])

    # 处理用户输入
    if prompt := st.chat_input("请输入您问题..."):
        # 立即显示用户消息
        with st.chat_message("user"):
            st.write(prompt)

        # 记录用户消息
        user_ts = datetime.now()
        st.session_state.user_message_timestamps.append(user_ts)
        st.session_state.messages.append({"role": "user", "content": prompt})

        try:
            # 构建带系统提示的消息 【关键修改】
            current_user_id = st.session_state.user_id
            messages = [
                {"role": "system", "content": SYSTEM_PROMPTS[current_user_id]},
                *st.session_state.messages[-5:]  # 保留最近5条对话
            ]

            start_time = time.time()
            response = requests.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers=headers,
                json={
                    "model": "deepseek-chat",
                    "messages": messages,
                    "temperature": 0.7
                }
            )
            response.raise_for_status()
            ai_reply = response.json()['choices'][0]['message']['content']

            # 计算API延迟
            api_latency = time.time() - start_time
            st.session_state.total_api_latency += api_latency

        except Exception as e:
            ai_reply = f"⚠️ 请求失败：{str(e)}"
            api_latency = 0

        # 立即显示AI回复
        with st.chat_message("assistant"):
            st.write(ai_reply)

        # 记录AI回复
        ai_ts = datetime.now()
        st.session_state.ai_response_timestamps.append(ai_ts)
        st.session_state.messages.append({"role": "assistant", "content": ai_reply})

        # 更新持续时间
        if st.session_state.user_message_timestamps and st.session_state.ai_response_timestamps:
            st.session_state.total_duration = (
                    st.session_state.ai_response_timestamps[-1] -
                    st.session_state.user_message_timestamps[0]
            ).total_seconds()
            st.session_state.net_duration = st.session_state.total_duration - st.session_state.total_api_latency

# ========== 以下部分保持不变 ==========
with st.container(height=400):
    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("最终成果")
        notes = st.text_area(
            "在此记录您的成果 👇",
            value=st.session_state.research_notes,
            height=300,
            key="notes_input"
        )
        st.session_state.research_notes = notes

    with col2:
        st.subheader("数据提交")
        questionnaire_url = "https://www.wjx.cn/vm/QHyr8WO.aspx"

        with st.form("submit_form"):
            submitted = st.form_submit_button("✅ 提交记录并跳转到量表填写部分")

            if submitted:
                if not notes.strip():
                    st.error("请先填写观察记录")
                else:
                    with open("research_notes.txt", "w", encoding="utf-8") as f:
                        f.write(f"实验时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                        f.write(f"用户ID：{st.session_state.user_id}\n")
                        f.write(f"总时长：{format_duration(st.session_state.total_duration)}\n")
                        f.write(f"净时长：{format_duration(st.session_state.net_duration)}\n")
                        f.write("\n观察记录：\n" + notes)

                    st.toast("记录已保存，正在跳转问卷页面...", icon="🚀")
                    st.session_state.should_redirect = True

        if st.session_state.get("should_redirect"):
            js_code = f"""
            <script>
                window.open('{questionnaire_url}', '_blank');
            </script>
            """
            html(js_code, height=0, width=0)
            st.session_state.should_redirect = False

            st.error(f"""
            ⚠️ 如果长时间未跳转，请直接访问：  
            [问卷链接]({questionnaire_url})  
            或手动复制：  
            `{questionnaire_url}`
            """)